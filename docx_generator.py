from docx import Document
from docx.shared import Inches

"""
document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
p.add_run('  Какой-то текст на Русском').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('/home/alex/PycharmProjects/NiceOligos/static/background_1.jpeg',
                     width=Inches(4.25))

records = [
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
]

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')

"""
rowdata = [
    {'#': '1', 'name': 'oligo 1', 'seq': 'ACGTACGTACGTACGTACGTACGTACGTACGT', 'OE': '10', 'nmol': '50', 'pyrif': 'RP-картридж' },
    {'#': '2', 'name': 'oligo 2', 'seq': 'ACGTACGTACGTACGTACGTACGTACGTACGT', 'OE': '11', 'nmol': '50', 'pyrif': 'RP-картридж' },
    {'#': '3', 'name': 'oligo 3', 'seq': 'ACGTACGTACGTACGTACGTACGTACGTACGT', 'OE': '12', 'nmol': '50', 'pyrif': 'RP-картридж' },
    {'#': '4', 'name': 'oligo 4', 'seq': 'ACGTACGTACGTACGTACGTACGTACGTACGT', 'OE': '11', 'nmol': '50', 'pyrif': 'RP-картридж' },
    {'#': '5', 'name': 'oligo 5', 'seq': 'ACGTACGTACGTACGTACGTACGTACGTACGT', 'OE': '9', 'nmol': '50', 'pyrif': 'RP-картридж' },
]

document = Document()
document.add_heading('ПАСПОРТ ПРОДУКТА (УТ-1428)', 2)

table = document.add_table(rows=1, cols=6)
hdr_cells = table.rows[0].cells
for i, key in enumerate(rowdata[0].keys()):
    hdr_cells[i].text = key

for row in rowdata:
    row_cells = table.add_row().cells
    for i, key in enumerate(row.keys()):
        row_cells[i].text = row[key]

document.add_page_break()

document.save('demo.docx')