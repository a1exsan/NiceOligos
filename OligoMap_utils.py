import pandas as pd
import requests
import json
from datetime import datetime
from nicegui import app, ui

from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer
from docx import Document

class api_db_interface():

    def __init__(self, db_IP, db_port):
        self.db_IP = db_IP
        self.db_port = db_port
        self.api_db_url = f'http://{self.db_IP}:{self.db_port}'
        self.pincode = ''
        self.client = {}
        self.client_frontend = {}
        self.db_name = 'scheduler_oligolab_2.db'
        self.lcms_db_name = 'lcms_data_1.db'
        self.mod_db_name = 'modification_base_1.db'
        self.synth_db_name = 'synth_method_1.db'
        self.widget_db = 'gui_object_content_1.db'
        self.stock_db_name = 'stock_oligolab_5.db'
        self.db_users = 'users_1.db'
        self.db_price_name = 'oligo_price_blmx_2.db'
        self.chrom_db = 'chrom_data_1.db'
        self.maps_db_name = 'asm2000_map_1.db'
        self.status_hist_db = 'oligo_status_history_1.db'

    def headers(self):
        return {'Authorization': f'Pincode {self.pincode}'}

    def check_pincode(self):
        url = f'{self.api_db_url}/check_pincode'
        ret = requests.get(url, headers=self.headers())
        return ret

    def get_user_data(self):
        url = f'{self.api_db_url}/auth'
        ret = requests.get(url, headers=self.headers())
        return ret

class docx_passport():
    def __init__(self, invoce_number, rowdata):
        self.invoce_number = invoce_number
        self.rowdata = rowdata

        self.path = 'templates'
        self.headers = [
            '№',
            'Название',
            '5’ – 3’ последовательность',
            'Кол-во ОЕ',
            'Кол-во нмоль',
            'V(H2O)100 мкМ р-р, мкл',
            'Очистка'
        ]
        self.column_widths_cm = [0.3, 1.4, 2.0, 0.7, 0.7, 0.7, 1.0]

    def set_cell_text(self, cell, text, font_name='Arial', font_size=Pt(10), bold=False):
        cell.text = ''
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        font = run.font
        font.name = font_name
        font.size = font_size
        font.bold = bold
        rFonts = run._element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), font_name)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def save_table(self):
        document = Document()
        table = document.add_table(rows=1, cols=7)
        for i in range(len(table.columns)):
            table.columns[i].width = Inches(self.column_widths_cm[i])
        hdr_cells = table.rows[0].cells
        for cell, text in zip(hdr_cells, self.headers):
            self.set_cell_text(cell, text, font_name='Times New Roman', font_size=Pt(12), bold=True)

        for row in self.rowdata:
            row_cells = table.add_row().cells
            self.set_cell_text(row_cells[0], str(row['#']), font_name='Calibri', font_size=Pt(8))
            self.set_cell_text(row_cells[1], str(row['Name']), font_name='Calibri', font_size=Pt(8))
            self.set_cell_text(row_cells[2], str(row['Sequence']), font_name='Calibri', font_size=Pt(8))
            self.set_cell_text(row_cells[3], str(row['Amount,_oe']), font_name='Calibri', font_size=Pt(8))
            self.set_cell_text(row_cells[4], str(row['Amount,_nmol']), font_name='Calibri', font_size=Pt(8))
            self.set_cell_text(row_cells[5], str(row['Desolving']), font_name='Calibri', font_size=Pt(8))
            self.set_cell_text(row_cells[6], str(row['Purification']), font_name='Calibri', font_size=Pt(8))
        table.style = 'TableGrid'
        document.save(f'{self.path}/passport_rowdata.docx')

    def compose_passport(self):
        self.save_table()
        master = Document(f"{self.path}/pass_tmpl.docx")
        composer = Composer(master)
        doc1 = Document(f"{self.path}/passport_rowdata.docx")
        doc2 = Document(f"{self.path}/pass_tmpl_2.docx")
        composer.append(doc1)
        composer.append(doc2)
        composer.save(f"{self.path}/combined.docx")
        doc = DocxTemplate(f"{self.path}/combined.docx")
        context = {
            'date': datetime.now().date().strftime('%d.%m.%Y'),
            'invoce_number': f'{self.invoce_number}',
        }
        doc.render(context)
        doc.save(f"{self.path}/passport_doc.docx")

class stock_write_inoff(api_db_interface):
    def __init__(self, rowdata):
        IP = app.storage.general.get('db_IP')
        port = app.storage.general.get('db_port')
        super().__init__(IP, port)
        self.pincode = app.storage.user.get('pincode')
        self.rowdata = rowdata
        self.strftime_format = "%Y-%m-%d"
        self.time_format = "%H:%M:%S"

    def get_all_data_in_tab_key(self, tab_name, key, value):
        url = f'{self.api_db_url}/get_keys_data/{self.stock_db_name}/{tab_name}/{key}/{value}'
        ret = requests.get(url, headers=self.headers())
        return ret.json()

    def substruct_solution(self, unicode, amount, tab_name):
        ret = self.get_all_data_in_tab_key('total_tab', 'unicode', unicode)
        data = json.loads(ret[0][4])
        if 'smart' in list(data.keys()):
            if data['smart']['mol_lumiprobe_data'] not in ['', '{}']:
                d_dict = json.loads(data['smart']['mol_lumiprobe_data'])
                for key in d_dict.keys():
                    amount_key = float(d_dict[key]) * float(amount)
                    if float(amount_key) > 0:
                        r_ret = self.get_all_data_in_tab_key('total_tab', 'unicode', key)
                        url = f"{self.api_db_url}/insert_data/{self.stock_db_name}/{tab_name}"
                        r = requests.post(url,
                                      json=json.dumps(
                                          [
                                              r_ret[0][1], key, amount_key,
                                              datetime.now().date().strftime(self.strftime_format),
                                              datetime.now().time().strftime(self.time_format),
                                              self.get_user_id()
                                          ]
                                      )
                                      , headers=self.headers())


    def substruct_from_stock(self, tab_name, rowdata):
        for row in rowdata:
            if row['write-off data'] > 0:
                #print('write-off')
                ret = self.get_all_data_in_tab_key('total_tab', 'unicode', row['unicode'])
                row['name'] = ret[0][1]
                if row['name'].find('_sol_') > -1:
                    self.substruct_solution(row['unicode'], row['write-off data'], tab_name)
                else:
                    try:
                        f_amount = float(row['write-off data'])
                    except:
                        f_amount = 0
                    if float(f_amount) > 0:
                        url = f"{self.api_db_url}/insert_data/{self.stock_db_name}/{tab_name}"
                        r = requests.post(url,
                        json=json.dumps(
                            [
                            row['name'], row['unicode'], row['write-off data'],
                            datetime.now().date().strftime(self.strftime_format),
                            datetime.now().time().strftime(self.time_format),
                            #user_id
                            self.get_user_id()
                            ]
                            )
                        , headers=self.headers())
                        if r.status_code == 200:
                            if tab_name == 'output_tab':
                                ui.notify('Реагенты успешно списаны')
                            else:
                                ui.notify('Реагенты успешно внесены')

    def write_off(self):
        self.substruct_from_stock('output_tab', self.rowdata)

    def write_in(self):
        self.substruct_from_stock('input_tab', self.rowdata)

    def get_user_id(self):
        url = f"{self.api_db_url}/get_keys_data/{self.db_users}/users/pass/{self.pincode}"
        r = requests.get(url, headers=self.headers())
        return r.json()[0][1]


class oligos_data_stack():
    def __init__(self):
        self.input_selected_rows = {}


class stock_db_data(api_db_interface):
    def __init__(self):
        IP = app.storage.general.get('db_IP')
        port = app.storage.general.get('db_port')
        super().__init__(IP, port)
        self.pincode = app.storage.user.get('pincode')
        self.db_name = self.stock_db_name#'stock_oligolab_5.db'
        #self.db_users = 'users_1.db'
        self.strftime_format = "%Y-%m-%d"
        self.time_format = "%H:%M:%S"

    def get_remaining_stock(self, unicode):
        url = f"{self.api_db_url}/get_remaining_stock/{self.db_name}/{unicode}"
        input_ret = requests.get(url, headers=self.headers())
        try:
            return input_ret.json()
        except:
            return {'exist': 0.}

    def get_all_data_in_tab(self, tab_name):
        url = f'{self.api_db_url}/get_all_tab_data/{self.db_name}/{tab_name}'
        ret = requests.get(url, headers=self.headers())
        if ret.status_code == 200:
            return ret.json()
        else:
            return []


    def get_all_users(self):
        url = f'{self.api_db_url}/get_all_tab_data/{self.db_users}/users'
        ret = requests.get(url, headers=self.headers())
        return ret.json()

    def get_total_tab_data(self):
        ret = self.get_all_data_in_tab('total_tab')
        if len(ret) > 0:
            rowdata = []
            for row in ret:
                d = {}
                d['#'] = row[0]
                d['Name'] = row[1]
                d['Unicode'] = row[2]
                try:
                    data_obj = json.loads(row[4])
                    d['Name'] = data_obj['smart']['name']
                    d['low limit'] = data_obj['smart']['low_limit']
                    d['units'] = data_obj['smart']['price_units']
                    d['producer'] = data_obj['smart']['producer']
                    d['supplyer'] = data_obj['smart']['supplyer']
                    if 'articul' in list(data_obj['smart'].keys()):
                        d['articul'] = data_obj['smart']['articul']
                    else:
                        d['articul'] = 'None'
                    d['price'] = float(data_obj['smart']['price'])
                except Exception as err:
                    d['low limit'] = row[5]
                    d['units'] = row[3]
                    d['producer'] = ''
                    d['supplyer'] = ''
                    d['price'] = 0.
                    #print(row[1], row[2])
                    #print(err)
                rowdata.append(d)
            return rowdata
        else:
            return []

    def get_in_out_tab(self, name):
        users = self.get_all_users()
        users2 = self.get_all_data_in_tab(f'users')
        ids = {}
        for user in users:
            ids[user[1]] = user[1]
        for user in users2:
            ids[user[2]] = user[1]
        data = self.get_all_data_in_tab(f'{name}')
        df = pd.DataFrame(data)
        tab_df = pd.DataFrame({
            '#': df[0],
            'Name': df[1],
            "Amount": df[3],
            'Unicode': df[2],
            "Date": df[4],
            "Time": df[5],
            'User': df[6]
        })
        tab_df['User'] = [ids[i] for i in tab_df['User']]
        return tab_df.to_dict('records')


class oligomaps_search(api_db_interface):

    def __init__(self, api_IP, db_port):
        super().__init__(api_IP, db_port)
        self.pincode = ''
        self.strftime_format = "%Y-%m-%d"
        #self.db_name = 'scheduler_oligolab_2.db'


    def map_in_progress(self, mapdata):
        map = json.loads(mapdata)
        df = pd.DataFrame(map)
        try:
            return round(df[df['Status'] == 'finished'].shape[0] * 100/df.shape[0], 0)
        except:
            return 100

    def check_chrom_data_in_base(self, map_id):
        url = f'{self.api_db_url}/check_chrom_data/{self.chrom_db}/{self.maps_db_name}/{map_id}'
        r = requests.get(url, headers=self.headers())
        return r

    def check_lcms_data_in_base(self, map_id):
        pos_list = []
        url = f"{self.api_db_url}/get_keys_data/{self.lcms_db_name}/tab/map_id/{map_id}"
        ret = requests.get(url, headers=self.headers())
        if ret.status_code == 200:
            for row in ret.json():
                pos_list.append(row[2])
        return pos_list

    def insert_chrom_data_to_base(self, data):
        if data != {}:
            url = f'{self.api_db_url}/insert_data/{self.chrom_db}/main_tab'
            oligo_id = data['oligo_id']
            map_id = data['map_id']
            position = data['position']
            date = datetime.now().date().strftime('%d.%m.%Y')
            format_d = '%d.%m.%Y'
            chrom_data = json.dumps(data['chrom_data'])
            r = requests.post(url,
                              json=json.dumps([
                                  oligo_id, map_id, position, date, format_d, chrom_data
                              ]), headers=self.headers())
            if r.status_code == 200:
                ui.notify('Хроматограмма добавлена в базу')
            return r.status_code == 200
        else:
            return False

    def insert_lcms_data_to_base(self, data):
        if data != {}:
            oligo_id = int(data['oligo_ID'])
            map_id = int(data['map_ID'])
            position = data['position']
            ID, data_ = self.get_lcms_data_id(map_id, position)
            date = datetime.now().date().strftime('%d.%m.%Y')
            format_d = '%d.%m.%Y'
            lcms_data = json.dumps(data)
            if ID == -1:
                url = f'{self.api_db_url}/insert_data/{self.lcms_db_name}/tab'
                r = requests.post(url,
                              json=json.dumps([
                                  map_id, position, oligo_id, date, format_d, lcms_data
                              ]), headers=self.headers())
                if r.status_code == 200:
                    ui.notify('данные добавлены в базу')
            else:
                url = f'{self.api_db_url}/update_data/{self.lcms_db_name}/tab/{ID}'
                #print(lcms_data)
                r = requests.put(url,
                                 json=json.dumps({
                                     'name_list': ['data_json', 'date'],
                                     'value_list': [lcms_data, date]
                                 }), headers=self.headers())
                if r.status_code == 200:
                    ui.notify('данные обновлены')
            return r.status_code == 200
        else:
            return False

    def get_lcms_data_id(self, map_id, position):
        ID = -1
        data = {}
        url = f"{self.api_db_url}/get_keys_data/{self.lcms_db_name}/tab/map_id/{map_id}"
        ret = requests.get(url, headers=self.headers())
        if ret.status_code == 200:
            for row in ret.json():
                if row[2] == position:
                    ID = row[0]
                    data = json.loads(row[6])
                    return ID, data
        else:
            ui.notify(f'не удается загрузить карту: {map_id}')
        return ID, data

    def load_lcms_data_from_base(self, map_id, position):
        ID, data = self.get_lcms_data_id(int(map_id), position)
        if ID > -1:
            return data
        else:
            return {}


    def update_chrom_data_to_base(self, data):
        if data != {}:
            url = f'{self.api_db_url}/update_data/{self.chrom_db}/main_tab/{data["ID"]}'
            chrom_data = json.dumps(data['chrom_data'])
            r = requests.put(url,
                              json=json.dumps({
                                'name_list':['chrom_data'],
                                'value_list':[chrom_data]
                              }), headers=self.headers())
            if r.status_code == 200:
                ui.notify('Хроматограмма обновлена')
            return r.status_code == 200
        else:
            return False

    def search_chrom_data(self, oligo_id, position):
        url = f'{self.api_db_url}/get_keys_data/{self.chrom_db}/main_tab/oligo_id/{oligo_id}'
        ret = requests.get(url, headers=self.headers())
        d = {}
        for row in ret.json():
            if row[3] == position:
                d['id'] = row[0]
                d['oligo_id'] = row[1]
                d['map_id'] = row[2]
                d['position'] = row[3]
                d['date'] = row[4]
                d['format'] = row[5]
                d['chrom_data'] = row[6]
                break
        return d

    def get_oligomaps(self):
        self.all_not_fin_oligos = []
        url = f'{self.api_db_url}/get_all_tab_data/{self.maps_db_name}/main_map'
        ret = requests.get(url, headers=self.headers())
        if ret.status_code == 200:
            out = []
            for r in ret.json():
                d = {}
                d['#'] = r[0]
                d['Map name'] = r[2]
                d['Synth number'] = r[3]
                d['Date'] = r[1]
                d['in progress'] = self.map_in_progress(r[4])
                #d['map data'] = pd.DataFrame(json.loads(r[4]))

                df = pd.DataFrame(json.loads(r[4]))
                if 'Status' in list(df.keys()):
                    self.all_not_fin_oligos.extend(df[df['Status'] != 'finished'].to_dict('records'))
                    d['Finished'] = df[df['Status'] == 'finished'].shape[0]
                else:
                    d['Finished'] = df.shape[0]
                d['Total'] = df.shape[0]
                if 'Wasted' in list(df.keys()):
                    w_df = df[df['Wasted'] == True]
                    d['Wasted'] = w_df.shape[0]
                else:
                    d['Wasted'] = 0
                out.append(d)
            return out
        else:
            return []

    def get_oligomaps_data(self):
        url = f'{self.api_db_url}/get_all_tab_data/{self.maps_db_name}/main_map'
        ret = requests.get(url, headers=self.headers())
        if ret.status_code == 200:
            out = []
            for r in ret.json():
                d = {}
                d['#'] = r[0]
                d['Map name'] = r[2]
                d['Synth number'] = r[3]
                d['Date'] = r[1]
                d['in progress'] = self.map_in_progress(r[4])
                d['map data'] = pd.DataFrame(json.loads(r[4]))
                d['accord data'] = pd.DataFrame(json.loads(r[5]))
                df = pd.DataFrame(d['map data'])
                if 'Status' in list(df.keys()):
                    d['Finished'] = df[df['Status'] == 'finished'].shape[0]
                else:
                    d['Finished'] = df.shape[0]
                d['Total'] = df.shape[0]
                if 'Wasted' in list(df.keys()):
                    d['Wasted'] = df[df['Wasted'] == True].shape[0]
                    #print(d['Wasted'])
                else:
                    d['Wasted'] = 0
                out.append(d)
            return out
        else:
            return []

    def get_actual_maps(self):
        total_maps = self.get_oligomaps_data()
        if len(total_maps) > 0:
            out = []
            for row in total_maps:
                df = row['map data']
                if df.shape[0] > 0:
                    if df[(df['DONE'] == True)|(df['Wasted'] == True)].shape[0] != df.shape[0]:
                        d = {}
                        d['#'] = row['#']
                        d['Map name'] = row['Map name']
                        d['Synth number'] = row['Synth number']
                        d['Date'] = row['Date']
                        d['in progress'] = row['in progress']
                        d['Finished'] = row['Finished']
                        d['Total'] = row['Total']
                        d['Wasted'] = row['Wasted']
                        out.append(d)
            return out
        else:
            return []

    def load_oligomap(self, seldata):
        if len(seldata) > 0:
            url = f"{self.api_db_url}/get_keys_data/{self.maps_db_name}/main_map/id/{seldata[0]['#']}"
            ret = requests.get(url, headers=self.headers())
            if ret.status_code == 200:
                self.oligo_map_id = seldata[0]['#']
                meta = ret.json()
                map_data = json.loads(meta[0][4])
                accord_data = json.loads(meta[0][5])
                #print(meta)
                return map_data, accord_data, meta[0][2], meta[0][3]
            else:
                return [], []

    def delete_map_from_base(self, seldata):
        if len(seldata) > 0:
            self.oligo_map_id = -1
            url = f"{self.api_db_url}/delete_data/{self.maps_db_name}/main_map/{seldata[0]['#']}"
            ret = requests.delete(url, headers=self.headers())
            return ret.status_code
        else:
            return 404

    def get_wasted_in_progress(self):
        actual_maps = self.get_actual_maps()
        out = []
        for map in actual_maps:
            map_data, accord_data, name, synth_num = self.load_oligomap([map])
            df = pd.DataFrame(map_data)
            d = [{'Order id': i} for i in df[df['Wasted'] == True]['Order id']]
            out.extend(d)
        return out

    def get_oligomaps_data_tail(self, tail_len):
        url = f'{self.api_db_url}/get_all_tab_data_tail/{self.maps_db_name}/main_map/{tail_len}'
        ret = requests.get(url, headers=self.headers())
        if ret.status_code == 200:
            out = []
            for r in ret.json():
                d = {}
                d['#'] = r[0]
                d['Map name'] = r[2]
                d['Synth number'] = r[3]
                d['Date'] = r[1]
                d['map data'] = json.loads(r[4])
                out.append(d)
            return out
        else:
            return []

    def get_oligomaps_date_tail(self, date, tail_len=20):
        url = f'{self.api_db_url}/get_oligomaps_date_tail/{self.maps_db_name}/main_map/{date}/{tail_len}'
        ret = requests.get(url, headers=self.headers())
        if ret.status_code == 200:
            out = []
            for r in ret.json():
                d = {}
                d['#'] = r['0']
                d['Map name'] = r['2']
                d['Synth number'] = r['3']
                d['Date'] = r['1']
                d['map data'] = json.loads(r['4'])
                out.append(d)
                #print(d['Synth number'])
            return out
        else:
            return []

    def find_amount_by_order_id(self, order_id):
        maps = pd.DataFrame()
        for map, syn in zip(self.map_list['map data'], self.map_list['Synth number']):
            map = pd.DataFrame(map)
            map['Synt number'] = syn
            maps = pd.concat([maps, map])
        maps.reset_index(inplace=True)
        maps = maps[(maps['Order id'] == int(order_id))&(maps['Wasted'] == False)]
        return maps

    def get_low_amount_limit(self, amount):
        if amount.find('-')>-1:
            return float(amount[:amount.find('-')])
        else:
            return float(amount)

    def insert_map_to_base(self, name, synth_number, rowData, accordData):
        if len(rowData) > 0:

            synth_date = datetime.now().date().strftime(self.strftime_format)

            url = f"{self.api_db_url}/insert_data/{self.maps_db_name}/main_map"
            r = requests.post(url,
                              json=json.dumps([synth_date, name, synth_number,
                                               json.dumps(rowData),
                                               json.dumps(accordData)]), headers=self.headers())

            url = f'{self.api_db_url}/get_all_tab_data/{self.maps_db_name}/main_map'
            ret = requests.get(url, headers=self.headers())
            indexs = [r[0] for r in ret.json()]

            self.oligo_map_id = max(indexs)

            data = []
            for row in rowData:
                d = row.copy()
                d['map #'] = self.oligo_map_id
                data.append(d)

            url = f"{self.api_db_url}/update_data/{self.maps_db_name}/main_map/{self.oligo_map_id}"
            r = requests.put(url,
                             json=json.dumps({
                                 'name_list': ['map_tab'],
                                 'value_list': [
                                     json.dumps(data)
                                 ]
                             })
                             , headers=self.headers())

            return r.status_code
        else:
            return 404


    def get_order_status(self, row):
        state_list = ['synth', 'sed', 'click', 'cart', 'hplc', 'paag', 'LCMS', 'subl']
        flag_list = []
        for state in state_list:
            flag_list.append(row[f'Do {state}'] == row[f'Done {state}'])
        status = 'synthesis'
        for i in range(8):
            if not flag_list[i]:
                if i < 3:
                    status = 'synthesis'
                elif i > 2 and i < 6:
                    status = 'purification'
                elif i == 7:
                    status = 'formulation'
                return status
            else:
                if i == 7:
                    status = 'finished'
                    return status
        return status


    def set_omap_status(self, rowdata):
        out = []
        for row in rowdata:
            d = row.copy()
            d['Status'] = self.get_order_status(row)
            out.append(d)
        return out

    def update_oligomap_status(self, rowData, accordrowdata):
        if len(rowData) > 0:
            #print(rowData[0])
            if 'map #' in list(rowData[0].keys()):
                for row in rowData:
                    if row['map #'] != '':
                        self.oligo_map_id = int(row['map #'])
                        print(f'MAP ID: {self.oligo_map_id}')
                        break
        if self.oligo_map_id > -1:
            out = []
            for row in rowData:
                out.append(row)
                out[-1]['Date'] = datetime.now().date().strftime('%d.%m.%Y')
                out[-1]['Status'] = self.get_order_status(row)
                if out[-1]['Status'] == 'finished':
                    out[-1]['DONE'] = True
                else:
                    out[-1]['DONE'] = False

            url = f"{self.api_db_url}/update_data/{self.maps_db_name}/main_map/{self.oligo_map_id}"
            r = requests.put(url,
                              json=json.dumps({
                                  'name_list': ['map_tab', 'accord_tab'],
                                  'value_list': [
                                      json.dumps(out),
                                      json.dumps(accordrowdata)
                                  ]
                              })
                             , headers=self.headers())
            print(f'update status {self.oligo_map_id}: {r.status_code}')
            return out
        else:
            return rowData


    def get_chenged_cells_list(self, old_rowdata, new_rowdata):
        df_old = pd.DataFrame(old_rowdata)
        df_old.fillna('none', inplace=True)
        df_new = pd.DataFrame(new_rowdata)
        df_new.fillna('none', inplace=True)

        try:
            mask = df_old != df_new
            changed_cells = [(idx, col) for idx, col in mask.stack()[lambda x: x].index]
            out_cells = []
            for cell in changed_cells:
                if df_new.at[cell] != df_old.at[cell]:
                    out_cells.append(cell)
            return out_cells
        except:
            out_cells = []
            for i in range(df_new.shape[0]):
                for col in df_new.keys():
                    out_cells.append((i, col))
            return out_cells


    def update_oligomap_order_status(self, rowData, accordrowdata, selrowdata):
        if len(rowData) > 0:
            if 'map #' in list(rowData[0].keys()):
                for row in rowData:
                    if row['map #'] != '':
                        self.oligo_map_id = int(row['map #'])
                        print(f'MAP ID: {self.oligo_map_id}')
                        break
        if self.oligo_map_id > -1:
            #print('HELLO', self.oligo_map_id)
            out = []
            for row in rowData:
                out.append(row)
                out[-1]['Date'] = datetime.now().date().strftime('%d.%m.%Y')
                out[-1]['Status'] = self.get_order_status(row)
                #print(out[-1]['Status'], out[-1]['Position'])
                if out[-1]['Status'] == 'finished':
                    out[-1]['DONE'] = True
                else:
                    out[-1]['DONE'] = False

            map_cells = self.get_chenged_cells_list(app.storage.user['init_map_rowdata'], out)
            accord_cells = self.get_chenged_cells_list(app.storage.user['init_accordtab_rowdata'], accordrowdata)

            #print('map_cells')
            #print(map_cells)

            url = f"{self.api_db_url}/update_oligomap/{self.maps_db_name}/{self.db_name}/main_map/{self.oligo_map_id}"
            r = requests.put(url,
                              json=json.dumps({
                                  'name_list': ['map_tab', 'accord_tab', 'selected', 'map_cells', 'accord_cells'],
                                  'value_list': [
                                      json.dumps(out),
                                      json.dumps(accordrowdata),
                                      json.dumps(selrowdata),
                                      json.dumps(map_cells),
                                      json.dumps(accord_cells)
                                  ]
                              })
                             , headers=self.headers())
            print(f'update status {self.oligo_map_id}: {r.status_code}')
            if r.status_code == 200:
                ui.notify('Данные обновлены')
            return out
        else:
            return rowData


    def update_order_status(self, rowData, selRowdata):
        df = pd.DataFrame(selRowdata)
        if len(rowData) > 0:
            for row in rowData:
                if df.shape[0] > 0:
                    if row['#'] in list(df['#']):
                        order_id = row['Order id']
                        order_date = row['Date']
                        order_status = row['Status']
                    else:
                        order_id = '###'
                else:
                    order_id = row['Order id']
                    order_date = row['Date']
                    order_status = row['Status']
                if order_id != '###':
                    url = f"{self.api_db_url}/update_data/{self.db_name}/orders_tab/{order_id}"
                    r = requests.put(url,
                    json=json.dumps({
                        'name_list': ['output_date', 'status'],
                        'value_list': [order_date, order_status]
                    })
                , headers=self.headers())



def test():
    osearch = oligomaps_search('127.0.0.1', '8012')
    osearch.pincode = '4720'
    #map_list = osearch.get_oligomaps_data_tail(40)
    #for row in map_list:
    #    print(row)

    map_list = osearch.get_oligomaps_date_tail('2025-07-02', 20)
    print(pd.DataFrame(map_list))

def test2():
    osearch = oligomaps_search('127.0.0.1', '8012')
    osearch.pincode = '4720'
    maps = osearch.find_amount_by_order_id(6205, '2025-07-02')
    print(maps.keys())
    df = maps['Dens, oe/ml'] * maps['Vol, ml']
    print(df.sum())

if __name__ == '__main__':
    test2()